#!/usr/bin/env python3
"""Build the cited India foreign policy parliamentary analysis as a .docx file."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys, os, re

sys.path.insert(0, os.path.dirname(__file__))

from cited_part1 import C1
from cited_part2 import C2
from cited_part3 import C3
from cited_part4 import C4
from cited_part5 import C5
from cited_footnotes import FOOTNOTES

FULL_TEXT = C1 + C2 + C3 + C4 + C5 + FOOTNOTES

# --- Superscript Unicode mapping ---
SUPER_DIGITS = {
    '⁰': '0', '¹': '1', '²': '2', '³': '3', '⁴': '4',
    '⁵': '5', '⁶': '6', '⁷': '7', '⁸': '8', '⁹': '9',
}
SUPER_CHARS = set(SUPER_DIGITS.keys())


def has_superscript(text):
    return any(ch in SUPER_CHARS for ch in text)


def split_with_superscripts(text):
    """Split text into alternating (normal, superscript) segment tuples."""
    segments = []
    current = []
    in_super = False
    for ch in text:
        if ch in SUPER_CHARS:
            if not in_super:
                if current:
                    segments.append(('normal', ''.join(current)))
                current = []
                in_super = True
            current.append(SUPER_DIGITS[ch])
        else:
            if in_super:
                if current:
                    segments.append(('super', ''.join(current)))
                current = []
                in_super = False
            current.append(ch)
    if current:
        segments.append(('super' if in_super else 'normal', ''.join(current)))
    return segments


def make_superscript(run):
    """Apply superscript vertical alignment to a run via XML."""
    rPr = run._r.get_or_add_rPr()
    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'), 'superscript')
    rPr.append(vertAlign)


# --- Document setup ---
doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.2)
    section.bottom_margin = Inches(1.2)
    section.left_margin   = Inches(1.3)
    section.right_margin  = Inches(1.3)


# --- Style helpers ---
def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(12)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_body_with_super(doc, text, indent=True):
    """Add a body paragraph, rendering superscript Unicode as real superscripts."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(15)

    if has_superscript(text):
        for kind, segment in split_with_superscripts(text):
            run = p.add_run(segment)
            run.font.size = Pt(11)
            if kind == 'super':
                make_superscript(run)
                run.font.size = Pt(7.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


def add_footnote_entry(doc, text):
    """Add a footnote entry line (smaller font, no indent)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Inches(0.3)
    p.paragraph_format.space_after   = Pt(4)
    p.paragraph_format.line_spacing  = Pt(13)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("━" * 60)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(12)
    return p


def add_page_break(doc):
    doc.add_page_break()


# --- Title page ---
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PARLIAMENT AS WITNESS AND ARCHITECT")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("The Evolution of India's Foreign Policy in Parliamentary Debates")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("1952 – 2025")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x8b, 0x00, 0x00)

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A Comprehensive Study Based on Lok Sabha and Rajya Sabha Proceedings")
run.italic = True
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Drawn from the Official Parliamentary Records of India\n"
                 "Covering the First Lok Sabha (1952) to the Eighteenth Lok Sabha (2025)\n"
                 "and the Rajya Sabha Continuous Record 1952–2025")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("With Citations and Footnotes — 430+ Parliamentary Sources")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
run.italic = True

add_page_break(doc)

# --- Parse and render content ---
in_footnotes_section = False

lines = FULL_TEXT.split('\n')
i = 0
while i < len(lines):
    line = lines[i].strip()

    if not line:
        i += 1
        continue

    # Divider lines
    if line.startswith('━'):
        add_divider(doc)
        i += 1
        continue

    # Detect transition to footnotes/notes section
    if 'NOTES AND CITATIONS' in line or 'SUPPLEMENTARY NOTES' in line:
        in_footnotes_section = True
        add_page_break(doc)
        add_heading1(doc, line)
        i += 1
        continue

    # SELECT BIBLIOGRAPHY
    if line.startswith('SELECT BIBLIOGRAPHY') or line == 'SELECT BIBLIOGRAPHY':
        add_page_break(doc)
        add_heading1(doc, line)
        i += 1
        continue

    # APPENDIX headings
    if re.match(r'^APPENDIX [A-Z]', line):
        add_page_break(doc)
        add_heading1(doc, line)
        i += 1
        continue

    # CHAPTER headings
    if line.startswith('CHAPTER ') and ':' in line:
        add_page_break(doc)
        add_heading1(doc, line)
        i += 1
        continue

    # All-caps headings (PREFACE, CONCLUSION, etc.)
    if line.isupper() and len(line) > 6 and not line.startswith('━'):
        add_heading1(doc, line)
        i += 1
        continue

    # Numbered subsections: 10.3 Title
    sec_match = re.match(r'^(\d+\.\d+)\s+(.+)$', line)
    if sec_match:
        add_heading3(doc, line)
        i += 1
        continue

    # Year-reference lines in appendix (1952, March: ...)
    year_match = re.match(r'^(\d{4})[,\s]', line)
    if year_match and ':' in line and len(line) < 300:
        p = doc.add_paragraph()
        parts = line.split(':', 1)
        run1 = p.add_run(parts[0] + ': ')
        run1.bold = True
        run1.font.size = Pt(10.5)
        if len(parts) > 1:
            rest = parts[1].strip()
            if has_superscript(rest):
                for kind, seg in split_with_superscripts(rest):
                    r = p.add_run(seg)
                    r.font.size = Pt(10.5)
                    if kind == 'super':
                        make_superscript(r)
                        r.font.size = Pt(7.5)
            else:
                r = p.add_run(rest)
                r.font.size = Pt(10.5)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Inches(0.2)
        i += 1
        continue

    # Footnote entries: "352. LSD, ..." (number followed by period at start)
    fn_match = re.match(r'^(\d+)\.\s+(.+)$', line)
    if fn_match and in_footnotes_section:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.line_spacing = Pt(13)
        num_run = p.add_run(fn_match.group(1) + '.  ')
        num_run.bold = True
        num_run.font.size = Pt(9.5)
        rest = fn_match.group(2)
        if has_superscript(rest):
            for kind, seg in split_with_superscripts(rest):
                r = p.add_run(seg)
                r.font.size = Pt(9.5)
                if kind == 'super':
                    make_superscript(r)
                    r.font.size = Pt(7)
        else:
            r = p.add_run(rest)
            r.font.size = Pt(9.5)
        i += 1
        continue

    # Regular body paragraph
    add_body_with_super(doc, line)
    i += 1

# --- End page ---
add_page_break(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— End of Study —")
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
word_count = len(FULL_TEXT.split())
run = p.add_run(f"Total word count: approximately {word_count:,} words")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output_path = "/home/user/ravirasputin.github.io/India_Foreign_Policy_Parliamentary_Analysis_1952_2025.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
print(f"Total words: {word_count:,}")
